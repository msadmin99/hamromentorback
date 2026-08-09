"""Word (.docx) import — walks paragraphs looking for a documented
question/option/answer pattern (numbered questions, lettered options,
"Answer:"/"Explanation:" markers). Preserves superscript/subscript (critical
for Physics/Chemistry notation — x², H₂O) and bold/italic by walking run
properties; extracts embedded images via the paragraph's own drawing
relationships. This is a documented-convention parser, not free-form NLP —
faculty following the template's structure get reliable results; documents
that don't follow the pattern degrade to "no questions found," not silent
corruption.

Subject, Chapter, Topic and Course are NOT part of the file — the admin
selects them for the whole batch on the Preview & Validate screen. Marks and
negative marks aren't part of the file either — those are set later in Exam
Management.

Expected structure (also shown in the downloadable .docx template):

    Q1. A ball is thrown vertically upward. What is its acceleration
    at the highest point?
    A) Zero
    B) g, downward
    C) g, upward
    D) Depends on mass
    Answer: B
    Explanation: At the highest point velocity is zero but gravity
    still acts downward.
"""
import re

from docx import Document
from docx.oxml.ns import qn

from .base import ParsedQuestion, save_temp_image

QUESTION_RE = re.compile(r'^\s*(?:Q\.?\s*)?(\d+)\s*[.):]\s*')
OPTION_RE = re.compile(r'^\s*([A-Da-d])\s*[.):]\s*')
ANSWER_RE = re.compile(r'^\s*Answer\s*:?\s*([A-Da-d])', re.IGNORECASE)
EXPLANATION_RE = re.compile(r'^\s*Explanation\s*:?\s*')


def _run_html(run, text):
    if not text:
        return ''
    escaped = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    if run.font.superscript:
        escaped = f'<sup>{escaped}</sup>'
    elif run.font.subscript:
        escaped = f'<sub>{escaped}</sub>'
    if run.italic:
        escaped = f'<em>{escaped}</em>'
    if run.bold:
        escaped = f'<strong>{escaped}</strong>'
    return escaped


def _paragraph_html(para, strip_prefix_len=0):
    parts = []
    consumed = 0
    for run in para.runs:
        text = run.text or ''
        if consumed < strip_prefix_len:
            cut = min(strip_prefix_len - consumed, len(text))
            text = text[cut:]
            consumed += cut
        if text:
            parts.append(_run_html(run, text))
    inline = ''.join(parts).strip()
    return f'<p>{inline}</p>' if inline else ''


def _paragraph_image(para, document, batch_id, slot):
    if not batch_id:
        return ''
    for run in para.runs:
        for blip in run._element.findall('.//' + qn('a:blip')):  # noqa: SLF001 - python-docx has no public run-image accessor
            rid = blip.get(qn('r:embed'))
            if not rid:
                continue
            try:
                blob = document.part.related_parts[rid].blob
            except KeyError:
                continue
            return save_temp_image(batch_id, slot, blob, ext='png')
    return ''


def parse_docx(file_obj, batch_id=None):
    try:
        document = Document(file_obj)
    except Exception as exc:  # noqa: BLE001 - surface any python-docx failure as a clean validation error
        raise ValueError(f'Could not read that Word document: {exc}') from exc

    questions = []
    current = None
    mode = None

    def flush():
        nonlocal current
        if current is not None and current['text_html']:
            questions.append(current)
        current = None

    for para in document.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue

        m = QUESTION_RE.match(raw_text)
        if m:
            flush()
            current = ParsedQuestion()
            mode = 'question'
            current['text_html'] = _paragraph_html(para, strip_prefix_len=m.end())
            img = _paragraph_image(para, document, batch_id, f'q{len(questions) + 1}_question')
            if img:
                current['question_image_path'] = img
            continue

        if current is None:
            continue  # stray text before the first recognized question marker

        m = OPTION_RE.match(raw_text)
        if m:
            mode = 'option'
            html_val = _paragraph_html(para, strip_prefix_len=m.end())
            img = _paragraph_image(para, document, batch_id, f'q{len(questions) + 1}_option{len(current["options"]) + 1}')
            current['options'].append({'text_html': html_val, 'is_correct': False, 'image_path': img})
            continue

        m = ANSWER_RE.match(raw_text)
        if m:
            letter = m.group(1).upper()
            for i, opt in enumerate(current['options']):
                opt['is_correct'] = chr(ord('A') + i) == letter
            mode = None
            continue

        m = EXPLANATION_RE.match(raw_text)
        if m and (raw_text.lower().startswith('explanation')):
            mode = 'explanation'
            current['explanation_html'] = _paragraph_html(para, strip_prefix_len=m.end())
            img = _paragraph_image(para, document, batch_id, f'q{len(questions) + 1}_explanation')
            if img:
                current['explanation_image_path'] = img
            continue

        # plain continuation of whichever section is currently open
        html_val = _paragraph_html(para)
        if not html_val:
            continue
        if mode == 'question':
            current['text_html'] += html_val
        elif mode == 'option' and current['options']:
            current['options'][-1]['text_html'] += html_val
        elif mode == 'explanation':
            current['explanation_html'] += html_val

    flush()
    return questions


def build_template_docx():
    document = Document()
    document.add_paragraph('Q1. A ball is thrown vertically upward. What is its acceleration at the highest point?')
    document.add_paragraph('A) Zero')
    document.add_paragraph('B) g, downward')
    document.add_paragraph('C) g, upward')
    document.add_paragraph('D) Depends on mass')
    document.add_paragraph('Answer: B')
    document.add_paragraph('Explanation: At the highest point velocity is zero but gravity still acts downward.')
    document.add_paragraph('')
    document.add_paragraph('Q2. Which ion is formed when sulfuric acid (H2SO4) fully dissociates, alongside SO4²⁻?')
    document.add_paragraph('A) H⁺')
    document.add_paragraph('B) OH⁻')
    document.add_paragraph('C) NH4⁺')
    document.add_paragraph('D) Fe³⁺')
    document.add_paragraph('Answer: A')
    document.add_paragraph('Explanation: H2SO4 -> 2H+ + SO4^2-')
    return document
